#!/usr/bin/env python3
"""
CUFA 카톡 zip docx 16개 → markdown 변환 + 메타 분석.

입력: 카톡 받은 파일에 풀어둔 docx
출력: scripts/research/cufa-docx/<file>.md + summary.json
실행: python scripts/research/extract-cufa-docx.py
"""

import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:  # pragma: no cover
    print("python-docx 미설치. `python -m pip install python-docx`", file=sys.stderr)
    sys.exit(1)

HERE = Path(__file__).resolve().parent
OUT_DIR = HERE / "cufa-docx"
OUT_DIR.mkdir(parents=True, exist_ok=True)

KAKAO = Path("C:/Users/lch68/Documents/카카오톡 받은 파일")

src_dirs = [
    KAKAO / "01_weekly_docx",
    KAKAO / "02_appendix_docx",
    Path("C:/Users/lch68/AppData/Local/Temp/cufa-zips/01_weekly/01_weekly_docx"),
    Path("C:/Users/lch68/AppData/Local/Temp/cufa-zips/02_appendix/02_appendix_docx"),
]

summary: list[dict] = []


def kind_of(src_dir: Path) -> str:
    s = str(src_dir).lower()
    return "weekly" if "weekly" in s else "appendix"


for src_dir in src_dirs:
    if not src_dir.exists():
        continue
    docx_files = sorted(src_dir.glob("*.docx"))
    if not docx_files:
        continue
    for docx_path in docx_files:
        # 중복 처리 (이미 동일 파일명 변환된 경우 스킵)
        out_path = OUT_DIR / (docx_path.stem + ".md")
        if out_path.exists() and any(s["file"] == docx_path.name for s in summary):
            continue

        doc = Document(str(docx_path))
        parts: list[str] = []
        headings: list[dict] = []

        for p in doc.paragraphs:
            t = p.text.strip()
            if not t:
                continue
            style = p.style.name if p.style else ""
            if style.startswith("Heading"):
                tail = style.split()[-1]
                level = int(tail) if tail.isdigit() else 1
                level = min(max(level, 1), 6)
                parts.append(f"{'#' * level} {t}")
                headings.append({"level": level, "text": t})
            elif style == "Title":
                parts.append(f"# {t}")
                headings.append({"level": 1, "text": t})
            else:
                parts.append(t)

        for ti, table in enumerate(doc.tables):
            parts.append(f"\n[table {ti+1}]")
            for row in table.rows:
                cells = [
                    cell.text.strip().replace("\n", " / ") for cell in row.cells
                ]
                parts.append("| " + " | ".join(cells) + " |")

        body = "\n\n".join(parts)
        out_path.write_text(body, encoding="utf-8")

        ko = len(re.findall(r"[\uac00-\ud7a3]", body))
        en = len(re.findall(r"[A-Za-z]", body))
        ko_ratio = ko / max(ko + en, 1)
        urls = len(re.findall(r"https?://[\w./%?=&#~+-]+", body))
        numbers = len(re.findall(r"\b\d{2,}\b", body))
        tables_n = len(doc.tables)
        chars = len(body)

        summary.append(
            {
                "file": docx_path.name,
                "kind": kind_of(src_dir),
                "out": out_path.name,
                "chars": chars,
                "headings_n": len(headings),
                "first_3_headings": [h["text"][:80] for h in headings[:3]],
                "tables": tables_n,
                "urls": urls,
                "numbers": numbers,
                "korean_ratio": round(ko_ratio, 3),
            }
        )
        print(
            f"v {docx_path.name:<14} -> {out_path.name:<14} chars={chars:>6} "
            f"headings={len(headings):>3} tables={tables_n:>2} urls={urls:>2}"
        )

(HERE / "cufa-docx-summary.json").write_text(
    json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
)
print(f"\n^ summary -> {HERE / 'cufa-docx-summary.json'} ({len(summary)} files)")
